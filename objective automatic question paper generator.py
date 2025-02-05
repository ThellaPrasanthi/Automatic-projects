from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import random
import re
import os
from datetime import datetime

def load_questions(input_file):
    # Load questions from the input .docx file
    doc = Document(input_file)
    questions = []
    question = ""
    for para in doc.paragraphs:
        if para.text.strip() != "":
            text = para.text.strip()
            if re.match(r'^\d+\.', text):  # Detect start of a new question
                if question:  # Append previous question
                    questions.append(question)
                question = text
            else:
                question += "\n" + text
    if question:  # Append the last question
        questions.append(question)
    return questions

def clean_question_text(text):
    # Clean the question text by removing the question number
    return re.sub(r'^\d+\.\s*', '', text)

def generate_question_paper(questions, num_questions):
    # Generate a question paper based on the user-defined number of questions
    if len(questions) < num_questions:
        raise ValueError("Not enough questions in the input file. Please add more questions.")
    return random.sample(questions, num_questions)

def create_header(doc, marks, branch, time_duration, exam_date, sub):
    # Create header for question paper
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run(
        f"SRI VENKATESWARA COLLEGE OF ENGINEERING, TIRUPATI\n"
        f"Name of the department: {branch}\n"
        f"Subject: {sub}\n"
        f"Marks: {marks}\n"
        f"Date: {exam_date}\t\tTime: {time_duration} mins\n\n"
        f"Answer the following questions"
    )
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run.font.size = Pt(12)

def save_question_paper(question_paper, output_file, marks, branch, time_duration, exam_date, sub):
    # Save the generated question paper to a .docx file
    doc = Document()
    
    # Add header
    create_header(doc, marks, branch, time_duration, exam_date, sub)
    
    for i, question in enumerate(question_paper, 1):
        # Process question and options
        parts = question.split('\n')
        question_text = clean_question_text(parts[0])
        options = [line.strip() for line in parts[1:] if line.strip()]
        
        # Add question
        paragraph = doc.add_paragraph(f"{i}. {question_text}           [   ]")
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Add options horizontally
        options_text = '    '.join(options)
        doc.add_paragraph(options_text)
    
    # Add footer information
    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "End of Question Paper"
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.save(output_file)

def main():
    input_file = input("Enter the path to the input .docx file: ")
    sub = input("Enter the exam subject: ")
    branch = input("Enter the name of the department: ")
    num_sets = int(input("Enter the number of question paper sets to generate: "))
    num_questions = int(input("Enter the number of questions per question paper: "))
    marks = input("Enter the total marks: ")
    time_duration = input("Enter the time duration: ")
    exam_date = input("Enter the date of the exam (dd-mm-yyyy): ")
    
    # Validate the date
    try:
        datetime.strptime(exam_date, "%d-%m-%Y")
    except ValueError:
        print("Invalid date format. Please enter the date in dd-mm-yyyy format.")
        return
    
    questions = load_questions(input_file)
    
    for i in range(num_sets):
        output_file = input(f"Enter the path to save Question Paper Set {i+1}: ")
        try:
            question_paper = generate_question_paper(questions, num_questions)
            save_question_paper(question_paper, output_file, marks, branch, time_duration, exam_date, sub)
        except ValueError as e:
            print(e)
            return
    
    print("Question papers generated successfully.")

if __name__ == "__main__":
    main()
